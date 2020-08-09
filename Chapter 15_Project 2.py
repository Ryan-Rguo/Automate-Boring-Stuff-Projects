# Chapter 15
# Practice Project 2

# Custom Invitations as Word Documents.

import docx, os

os.chdir('d:\\a\\b\\cats')

def invi():
     doc = docx.Document('demo.docx')
     txt = open('guests.txt','r')
     guests = txt.readlines()
     for i in range(len(guests)):
          name = guests[i].rstrip('\n')
          if i != 0:
               doc.add_paragraph('','python_title')
          doc.add_paragraph('It would be a pleasure to have the company of','python_title')
          doc.add_paragraph(name,'python_company')
          doc.add_paragraph('at 11010 Memory Lane on the Evening of','python_title')
          doc.add_paragraph('April 1st','python_date')
          doc.add_paragraph("at 7 o'clock",'python_title')
          if i + 1 != len(guests):
               doc.paragraphs[5 + (6 * i)].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
     doc.save('GuestInvi.docx')
     print('Done')

